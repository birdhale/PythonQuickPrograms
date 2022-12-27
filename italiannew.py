import requests
from bs4 import BeautifulSoup
from docx import Document
import random
import openpyxl

# Set the URL for the news website and send a request to the website (Change the URL to your Websites)

URL = "https://easyitaliannews.com/"
page = requests.get(URL)

# Parse the HTML content of the page

soup = BeautifulSoup(page.content, "html.parser")

# Find all the paragraphs on the website

paragraphs = soup.find_all("p")

# Select a random paragraph

if paragraphs:
    selected_paragraph = random.choice(paragraphs)
else:
    selected_paragraph = "Unable to find any paragraphs on the website"

# Open the Word document - Select your word document here

document = Document("random_paragraph.docx")

# Add the selected paragraph to the document

document.add_paragraph(selected_paragraph.text)

# Extract all the words from the document

words = []
for paragraph in document.paragraphs:
    for run in paragraph.runs:
        words.extend(run.text.split())

# Select 10 random words from the list of words

if len(words) >= 10:
    selected_words = random.sample(words, 10)
else:
    selected_words = words

# Open the Excel file / Select your excel file here

workbook = openpyxl.load_workbook("existing_file.xlsx")

# Get the first sheet in the workbook

sheet = workbook.active

# Find the next available row in the sheet

next_row = sheet.max_row + 1

# Paste the selected words into the Excel file, with each word in a different row

for i, word in enumerate(selected_words):
    sheet.cell(row=next_row + i, column=1).value = word

# Save the Excel file

workbook.save("existing_file.xlsx")

# Save the Word document

document.save("random_paragraph.docx")
